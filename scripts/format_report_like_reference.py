from copy import deepcopy
from pathlib import Path
import re
import shutil
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT / "final_report_completed_enriched.docx"
OUTPUT = ROOT / "期末报告_套用final_report_completed_enriched格式.docx"


def find_source() -> Path:
    downloads = Path.home() / "Downloads"
    matches = [
        p
        for p in downloads.glob("*.docx")
        if "2025-2026" in p.name and "软件测试与质量保证期末报告" in p.name and "第N组" in p.name
    ]
    if not matches:
        raise FileNotFoundError("Could not find the source report in Downloads.")
    return max(matches, key=lambda p: p.stat().st_mtime)


def set_east_asia_font(run_or_style, font_name: str) -> None:
    rpr = run_or_style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name: str, size_pt: float | None = None, bold=None) -> None:
    style.font.name = font_name
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    set_east_asia_font(style, font_name)


def copy_package_parts(src_docx: Path, dst_docx: Path, part_names: list[str]) -> None:
    tmp = dst_docx.with_suffix(".tmp.docx")
    with zipfile.ZipFile(dst_docx, "r") as zin, zipfile.ZipFile(src_docx, "r") as zref, zipfile.ZipFile(
        tmp, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        ref_names = set(zref.namelist())
        for item in zin.infolist():
            data = zref.read(item.filename) if item.filename in part_names and item.filename in ref_names else zin.read(item.filename)
            zout.writestr(item, data)
    tmp.replace(dst_docx)


def copy_paragraph_format(dst_p, src_p) -> None:
    dst_p.style = src_p.style.name
    dst_p.alignment = src_p.alignment
    dst_pPr = dst_p._p.get_or_add_pPr()
    src_pPr = src_p._p.pPr
    for child in list(dst_pPr):
        dst_pPr.remove(child)
    if src_pPr is not None:
        for child in src_pPr:
            dst_pPr.append(deepcopy(child))


def set_run_text_style(run, font_name="宋体", size=10.5, bold=None, italic=None) -> None:
    run.font.name = font_name
    set_east_asia_font(run, font_name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def is_chapter(text: str) -> bool:
    return bool(re.match(r"^第[一二三四五六七八九十]+章\s*", text.strip()))


def heading_level(text: str) -> int | None:
    stripped = text.strip()
    if is_chapter(stripped):
        return 1
    if re.match(r"^\d+\.\d+\.\d+\s*", stripped):
        return 3
    if re.match(r"^\d+\.\d+\s*", stripped):
        return 2
    return None


def format_paragraphs(doc: Document, ref: Document) -> None:
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if idx <= 44 and idx < len(ref.paragraphs):
            copy_paragraph_format(para, ref.paragraphs[idx])
            for run in para.runs:
                if text == "目  录":
                    set_run_text_style(run, "黑体", 16)
                elif idx in (1, 2):
                    set_run_text_style(run, "宋体", 36, bold=True)
                elif idx == 8:
                    set_run_text_style(run, "黑体", 18)
                elif para.style.name.startswith("toc"):
                    set_run_text_style(run, "宋体", 12)
            continue

        level = heading_level(text)
        if level == 1:
            para.style = "Heading 1"
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = None
            para.paragraph_format.line_spacing = 1.3
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                set_run_text_style(run, "黑体", 15.5, bold=True)
        elif level == 2:
            para.style = "Heading 2"
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = None
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                set_run_text_style(run, "楷体", 15, bold=False)
        elif level == 3:
            para.style = "Heading 3"
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = None
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                set_run_text_style(run, "楷体", 14, bold=True)
        else:
            para.style = "Normal"
            if text:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = Pt(21)
                para.paragraph_format.line_spacing = 1.08
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                set_run_text_style(run, "宋体", 10.5)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, left=120, bottom=80, right=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")


def set_table_borders(table, color="7F7F7F", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def format_cell_text(cell, font="宋体", size=10.5, bold=False, align=None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        for r in p.runs:
            set_run_text_style(r, font, size, bold=bold)


def format_tables(doc: Document) -> None:
    for ti, table in enumerate(doc.tables):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        table.autofit = False
        set_table_borders(table)

        if ti == 0:
            widths = [1425, 4645]
            set_table_width(table, sum(widths))
            for row in table.rows:
                for ci, cell in enumerate(row.cells):
                    set_cell_width(cell, widths[min(ci, 1)])
                    set_cell_margins(cell, 80, 100, 80, 100)
                    format_cell_text(
                        cell,
                        font="宋体",
                        size=12 if ci == 0 else 14,
                        bold=False,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                    )
        elif ti == len(doc.tables) - 1 and len(table.rows) == 1 and len(table.columns) == 1:
            set_table_width(table, 8200)
            cell = table.cell(0, 0)
            set_cell_width(cell, 8200)
            set_cell_margins(cell, 220, 220, 220, 220)
            format_cell_text(cell, font="宋体", size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
        else:
            cols = len(table.columns)
            total = 8200
            widths = [total // cols] * cols
            widths[-1] += total - sum(widths)
            set_table_width(table, total)
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    set_cell_width(cell, widths[ci])
                    set_cell_margins(cell)
                    set_cell_shading(cell, "EDEDED" if ri == 0 else "FFFFFF")
                    format_cell_text(
                        cell,
                        font="宋体",
                        size=10.5,
                        bold=ri == 0,
                        align=WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT,
                    )


def apply_reference_sections(doc: Document, ref: Document) -> None:
    ref_section = ref.sections[0]
    for section in doc.sections:
        section.page_width = ref_section.page_width
        section.page_height = ref_section.page_height
        section.top_margin = ref_section.top_margin
        section.bottom_margin = ref_section.bottom_margin
        section.left_margin = ref_section.left_margin
        section.right_margin = ref_section.right_margin
        section.header_distance = ref_section.header_distance
        section.footer_distance = ref_section.footer_distance


def format_styles(doc: Document) -> None:
    set_style_font(doc.styles["Normal"], "宋体", 10.5)
    set_style_font(doc.styles["Heading 1"], "黑体", 15.5, bold=True)
    set_style_font(doc.styles["Heading 2"], "楷体", 13, bold=True)
    set_style_font(doc.styles["Heading 3"], "楷体", 11, bold=True)


def main() -> None:
    source = find_source()
    shutil.copy2(source, OUTPUT)

    doc = Document(str(OUTPUT))
    ref = Document(str(REFERENCE))

    apply_reference_sections(doc, ref)
    format_styles(doc)
    format_paragraphs(doc, ref)
    format_tables(doc)

    doc.save(str(OUTPUT))
    copy_package_parts(
        REFERENCE,
        OUTPUT,
        [
            "word/header1.xml",
            "word/footer1.xml",
            "word/footer2.xml",
            "word/footer3.xml",
            "word/theme/theme1.xml",
            "word/fontTable.xml",
        ],
    )
    print(OUTPUT)


if __name__ == "__main__":
    main()
